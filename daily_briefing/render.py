import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def render_report(
    articles: list[dict],
    top_article: dict,
    insight_text: str,
    date_str: str,
) -> str:
    doc = Document()

    # Title
    title = doc.add_heading(f"Daily Briefing — {date_str}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("오전 9시 기준 | EV/Battery Industry")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Articles table
    _add_heading(doc, "오늘의 주요 기사", 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["#", "카테고리", "제목 / 요약", "링크"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True

    for a in articles:
        row = table.add_row().cells
        row[0].text = str(a["id"])
        row[1].text = a.get("category", "")
        summary_lines = a.get("summary", "").replace("\\n", "\n")
        row[2].text = f'{a["title"]}\n{summary_lines}'
        row[3].text = a.get("url", "")

    doc.add_paragraph()

    # Insight section
    _add_heading(doc, f'주목 기사: {top_article["title"]}', 1)
    doc.add_paragraph(f'카테고리: {top_article.get("category", "")} | Impact Score: {top_article.get("impact_score", "")}')
    doc.add_paragraph()

    for line in insight_text.split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(4)

    # Save
    os.makedirs("reports", exist_ok=True)
    path = os.path.join("reports", f"{date_str}_briefing.docx")
    doc.save(path)
    return path
